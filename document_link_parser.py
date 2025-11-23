"""
Document Link Parser for Groww Mutual Fund Links
Extracts links from .docx, .txt, .csv, and other document formats
"""

import re
from typing import List, Dict, Optional
from urllib.parse import urlparse
from pathlib import Path


class DocumentLinkParser:
    def __init__(self):
        self.base_url = "https://groww.in"
    
    def validate_groww_url(self, url: str) -> bool:
        """Validate that URL is a Groww mutual fund URL"""
        if not url or not isinstance(url, str):
            return False
        
        parsed = urlparse(url)
        if parsed.netloc not in ['groww.in', 'www.groww.in']:
            return False
        
        # Check if it's a mutual fund URL
        if '/mutual-funds/' not in url:
            return False
        
        return True
    
    def extract_urls_from_text(self, text: str) -> List[str]:
        """Extract all URLs from text"""
        # Pattern to match URLs
        url_pattern = re.compile(
            r'http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\\(\\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+',
            re.IGNORECASE
        )
        
        urls = url_pattern.findall(text)
        # Clean up URLs (remove trailing punctuation)
        cleaned_urls = []
        for url in urls:
            # Remove common trailing characters
            url = url.rstrip('.,;:!?)')
            if self.validate_groww_url(url):
                cleaned_urls.append(url)
        
        return cleaned_urls
    
    def parse_docx(self, file_path: str) -> Dict[str, List[str]]:
        """Parse .docx file and extract links organized by AMC"""
        try:
            from docx import Document
        except ImportError:
            raise ImportError(
                "python-docx is required for .docx files. Install it with: pip install python-docx"
            )
        
        doc = Document(file_path)
        all_urls = []
        amc_schemes = {}
        current_amc = None
        
        # Extract text from all paragraphs
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            
            # Extract URLs from paragraph text
            urls = self.extract_urls_from_text(text)
            all_urls.extend(urls)
            
            # Try to identify AMC names (usually in bold or as headings)
            # Look for patterns like "AMC Name:" or just check if it's a heading
            if paragraph.style.name.startswith('Heading') or not urls:
                # Might be an AMC name
                if text and len(text) < 100 and not text.startswith('http'):
                    # Could be an AMC name
                    if current_amc is None or urls:
                        current_amc = text
            
            # If we found URLs, associate them with current AMC
            if urls and current_amc:
                if current_amc not in amc_schemes:
                    amc_schemes[current_amc] = []
                amc_schemes[current_amc].extend(urls)
        
        # Also check hyperlinks in the document
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if run.hyperlink and run.hyperlink.address:
                    url = run.hyperlink.address
                    if self.validate_groww_url(url):
                        if url not in all_urls:
                            all_urls.append(url)
                            if current_amc:
                                if current_amc not in amc_schemes:
                                    amc_schemes[current_amc] = []
                                if url not in amc_schemes[current_amc]:
                                    amc_schemes[current_amc].append(url)
        
        # If we couldn't organize by AMC, return all URLs
        if not amc_schemes:
            return {
                'all_schemes': all_urls,
                'organized_by_amc': {}
            }
        
        return {
            'all_schemes': all_urls,
            'organized_by_amc': amc_schemes
        }
    
    def parse_txt(self, file_path: str) -> Dict[str, List[str]]:
        """Parse .txt file and extract links"""
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
        
        urls = self.extract_urls_from_text(text)
        
        # Try to organize by AMC (look for patterns)
        amc_schemes = {}
        lines = text.split('\n')
        current_amc = None
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Check if line contains URLs
            line_urls = self.extract_urls_from_text(line)
            
            # Check if line might be an AMC name (no URL, short, might be in caps or have special formatting)
            if not line_urls and len(line) < 100 and not line.startswith('http'):
                # Could be AMC name
                if any(keyword in line.lower() for keyword in ['mutual fund', 'amc', 'fund', 'capital']):
                    current_amc = line
                elif current_amc is None:
                    current_amc = line
            
            if line_urls:
                if current_amc:
                    if current_amc not in amc_schemes:
                        amc_schemes[current_amc] = []
                    amc_schemes[current_amc].extend(line_urls)
        
        return {
            'all_schemes': urls,
            'organized_by_amc': amc_schemes if amc_schemes else {}
        }
    
    def parse_csv(self, file_path: str) -> Dict[str, List[str]]:
        """Parse .csv file and extract links"""
        import csv
        
        all_urls = []
        amc_schemes = {}
        
        with open(file_path, 'r', encoding='utf-8') as f:
            reader = csv.reader(f)
            headers = next(reader, None)
            
            # Try to find AMC and URL columns
            amc_col = None
            url_col = None
            
            if headers:
                for i, header in enumerate(headers):
                    if 'amc' in header.lower() or 'fund house' in header.lower():
                        amc_col = i
                    if 'url' in header.lower() or 'link' in header.lower():
                        url_col = i
            
            for row in reader:
                if url_col is not None and url_col < len(row):
                    url = row[url_col].strip()
                    if self.validate_groww_url(url):
                        all_urls.append(url)
                        
                        # Organize by AMC if available
                        if amc_col is not None and amc_col < len(row):
                            amc = row[amc_col].strip()
                            if amc:
                                if amc not in amc_schemes:
                                    amc_schemes[amc] = []
                                amc_schemes[amc].append(url)
                else:
                    # No URL column, extract from all cells
                    for cell in row:
                        urls = self.extract_urls_from_text(cell)
                        all_urls.extend(urls)
        
        return {
            'all_schemes': list(set(all_urls)),
            'organized_by_amc': amc_schemes
        }
    
    def parse_document(self, file_path: str) -> Dict[str, List[str]]:
        """Parse document based on file extension"""
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        extension = path.suffix.lower()
        
        if extension == '.docx':
            return self.parse_docx(file_path)
        elif extension == '.txt':
            return self.parse_txt(file_path)
        elif extension == '.csv':
            return self.parse_csv(file_path)
        else:
            raise ValueError(f"Unsupported file format: {extension}. Supported: .docx, .txt, .csv")
    
    def save_extracted_links(self, links_data: Dict[str, List[str]], output_file: str = 'extracted_links.json'):
        """Save extracted links to JSON file"""
        import json
        
        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump(links_data, f, indent=2, ensure_ascii=False)
        
        print(f"Extracted links saved to {output_file}")
        print(f"Total schemes found: {len(links_data.get('all_schemes', []))}")
        print(f"AMCs found: {len(links_data.get('organized_by_amc', {}))}")


def main():
    """Example usage"""
    parser = DocumentLinkParser()
    
    # Update this path to your document
    doc_path = r"D:\download\Groww links for mutual funds.docx"
    
    try:
        print(f"Parsing document: {doc_path}")
        links_data = parser.parse_document(doc_path)
        
        print(f"\n✅ Successfully extracted links!")
        print(f"Total schemes: {len(links_data['all_schemes'])}")
        
        if links_data['organized_by_amc']:
            print(f"\nOrganized by AMC:")
            for amc, schemes in links_data['organized_by_amc'].items():
                print(f"  {amc}: {len(schemes)} schemes")
        
        # Save to JSON
        parser.save_extracted_links(links_data, 'groww_all_scheme_links.json')
        
        # Print first few URLs
        print(f"\nFirst 5 URLs:")
        for i, url in enumerate(links_data['all_schemes'][:5], 1):
            print(f"  {i}. {url}")
        
    except FileNotFoundError as e:
        print(f"❌ Error: {e}")
        print(f"Please update the file path in the script")
    except ImportError as e:
        print(f"❌ Error: {e}")
    except Exception as e:
        print(f"❌ Error parsing document: {e}")


if __name__ == "__main__":
    main()

